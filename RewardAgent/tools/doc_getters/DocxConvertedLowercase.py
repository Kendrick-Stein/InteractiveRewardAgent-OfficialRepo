from __future__ import annotations

import json
import os
from typing import List, Tuple

from smolagents import Tool
from docx import Document


def _find_uppercase_runs(doc: Document) -> List[Tuple[str, int]]:
    """
    Scan paragraphs and tables. If any run.text is uppercase (per str.isupper()),
    collect a small sample (text snippet, context type) for reporting.
    Returns a list of tuples (snippet, context), where context is 0 for paragraph, 1 for table cell.
    """
    findings: List[Tuple[str, int]] = []

    # Paragraphs
    for p in doc.paragraphs:
        for r in p.runs:
            t = r.text or ""
            if t and t.strip() and t.isupper():
                findings.append((t.strip()[:50], 0))
                if len(findings) >= 10:
                    return findings

    # Tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        t = r.text or ""
                        if t and t.strip() and t.isupper():
                            findings.append((t.strip()[:50], 1))
                            if len(findings) >= 10:
                                return findings
    return findings


class DocxConvertedLowercaseTool(Tool):
    name = "docx_converted_lowercase"
    description = (
        "Verify that all uppercase text in the DOCX has been converted to lowercase (tables and paragraphs). "
        "IMPORTANT: file_path must be a Host path. If your file is inside the VM (e.g., /home/user/...), "
        "use get_vm_file(vm_path, dest_name) to download it to Host first and pass the returned Host path."
    )
    inputs = {
        "file_path": {
            "description": "Path to the .docx file on the Host (absolute or relative).",
            "type": "string",
        }
    }
    output_type = "string"

    def forward(self, file_path: str) -> str:
        return self.__call__(file_path)

    def __call__(self, file_path: str) -> str:
        try:
            if not file_path:
                return json.dumps({"passed": False, "details": "file_path is required"}, ensure_ascii=False)
            abs_path = os.path.abspath(file_path)
            if not os.path.exists(abs_path):
                return json.dumps({
                    "passed": False,
                    "details": (
                        f"File not found on host: {abs_path}. If your file is in the VM, use get_vm_file to download it first."
                    )
                }, ensure_ascii=False)

            doc = Document(abs_path)
            findings = _find_uppercase_runs(doc)
            if findings:
                samples = [f"PARA: '{t}'" if c == 0 else f"TABLE: '{t}'" for t, c in findings[:5]]
                return json.dumps({
                    "passed": False,
                    "details": f"Found uppercase runs not converted to lowercase. Examples: {samples}"
                }, ensure_ascii=False)
            else:
                return json.dumps({
                    "passed": True,
                    "details": "All text appears converted to lowercase (no fully uppercase runs found)"
                }, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"passed": False, "details": f"Error: {e}"}, ensure_ascii=False)

    def to_code_prompt(self) -> str:
        return (
            "def docx_converted_lowercase(file_path: str) -> str:\n"
            "    '''Verify that uppercase text has been converted to lowercase (tables and paragraphs).\n"
            "    file_path must be a Host path; if the file is in the VM, first use get_vm_file.\n"
            "    Returns a JSON string with keys passed (bool) and details (str).'''\n"
        )
