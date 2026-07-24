from __future__ import annotations

import json
import os

from smolagents import Tool


class HasPageNumbersInFootersTool(Tool):
    name = "has_page_numbers_in_footers"
    description = (
        "Check whether a DOCX file has numeric page numbers in the footer across sections. "
        "IMPORTANT: file_path must be a Host path. If your file is inside the VM (e.g., /home/user/...), "
        "use get_vm_file(vm_path, dest_name) to download it to Host first and pass the returned Host path."
    )
    inputs = {
        "file_path": {
            "description": "Path to the .docx file on the Host (absolute or relative).",
            "type": "string",
        }
    }
    output_type = "string"

    def forward(self, file_path: str) -> str:
        return self.__call__(file_path)

    def __call__(self, file_path: str) -> str:
        try:
            if not file_path:
                return json.dumps({"passed": False, "details": "file_path is required"}, ensure_ascii=False)
            abs_path = os.path.abspath(file_path)
            if not os.path.exists(abs_path):
                return json.dumps({
                    "passed": False,
                    "details": (
                        f"File not found on host: {abs_path}. If your file is in the VM, use get_vm_file to download it first."
                    )
                }, ensure_ascii=False)

            try:
                from docx import Document
            except Exception as e:
                return json.dumps({"passed": False, "details": f"python-docx import error: {e}"}, ensure_ascii=False)

            try:
                doc = Document(abs_path)
            except Exception as e:
                return json.dumps({"passed": False, "details": f"Failed to open DOCX: {e}"}, ensure_ascii=False)

            # Check footer paragraphs across all sections and ensure at least one digit exists in each footer
            for section in doc.sections:
                footer = section.footer
                if footer is None:
                    return json.dumps({"passed": False, "details": "A section has no footer"}, ensure_ascii=False)
                footer_text = footer.paragraphs[0].text if footer.paragraphs else ''
                if not any(ch.isdigit() for ch in footer_text):
                    return json.dumps({"passed": False, "details": "No page number detected in a footer"}, ensure_ascii=False)

            return json.dumps({"passed": True, "details": "Page numbers detected in footers"}, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"passed": False, "details": f"Error: {e}"}, ensure_ascii=False)

    def to_code_prompt(self) -> str:
        return (
            "def has_page_numbers_in_footers(file_path: str) -> str:\n"
            "    '''Check whether a DOCX has numeric page numbers in its footers.\n"
            "    file_path must be a Host path; if the file is in the VM, first use get_vm_file.\n"
            "    Returns a JSON string with keys passed (bool) and details (str).'''\n"
        )
