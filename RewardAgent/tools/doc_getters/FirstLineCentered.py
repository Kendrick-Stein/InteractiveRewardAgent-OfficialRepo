from __future__ import annotations

import json
import os

from smolagents import Tool

# Use existing implementation from desktop_env
from desktop_env.evaluators.metrics.docs import is_first_line_centered as _is_first_line_centered


class FirstLineCenteredTool(Tool):
    name = "first_line_centered"
    description = (
        "Check whether the first paragraph of a DOCX file is center aligned. "
        "IMPORTANT: file_path must be a Host path. If your file is inside the VM (e.g., /home/user/...), "
        "use get_vm_file(vm_path, dest_name) to download it to Host first and pass the returned Host path."
    )
    inputs = {
        "file_path": {
            "description": "Path to the .docx file on the Host (absolute or relative).",
            "type": "string",
        }
    }
    output_type = "string"

    def forward(self, file_path: str) -> str:
        return self.__call__(file_path)

    def __call__(self, file_path: str) -> str:
        try:
            if not file_path:
                return json.dumps({"passed": False, "details": "file_path is required"}, ensure_ascii=False)
            abs_path = os.path.abspath(file_path)
            if not os.path.exists(abs_path):
                return json.dumps({
                    "passed": False,
                    "details": (
                        f"File not found on host: {abs_path}. If your file is in the VM, use get_vm_file to download it first."
                    )
                }, ensure_ascii=False)

            try:
                from docx import Document
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            except Exception as e:
                return json.dumps({"passed": False, "details": f"python-docx import error: {e}"}, ensure_ascii=False)

            try:
                doc = Document(abs_path)
            except Exception as e:
                return json.dumps({"passed": False, "details": f"Failed to open DOCX: {e}"}, ensure_ascii=False)

            if not doc.paragraphs:
                return json.dumps({"passed": False, "details": "Document has no paragraphs"}, ensure_ascii=False)

            passed = (doc.paragraphs[0].paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
            return json.dumps({
                "passed": bool(passed),
                "details": "First line is centered" if passed else "First line is not centered"
            }, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"passed": False, "details": f"Error: {e}"}, ensure_ascii=False)

    def to_code_prompt(self) -> str:
        return (
            "def first_line_centered(file_path: str) -> str:\n"
            "    '''Check whether the first paragraph is center aligned.\n"
            "    file_path must be a Host path; if the file is in the VM, first use get_vm_file.\n"
            "    Returns a JSON string with keys passed (bool) and details (str).'''\n"
        )
