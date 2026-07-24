from __future__ import annotations

import json
import os
from typing import Optional

from smolagents import Tool
from docx import Document
from docx.enum.text import WD_LINE_SPACING


def _get_spacing_value(p) -> Optional[float]:
    """
    Try to normalize the paragraph line spacing into a float multiplier.
    Returns 1.0, 1.5, 2.0 for SINGLE/ONE_POINT_FIVE/DOUBLE where possible,
    or the numeric value when rule is MULTIPLE, otherwise None if not determinable.
    """
    pf = p.paragraph_format
    rule = getattr(pf, "line_spacing_rule", None)
    ls = getattr(pf, "line_spacing", None)

    try:
        if rule == WD_LINE_SPACING.SINGLE:
            return 1.0
        if rule == WD_LINE_SPACING.ONE_POINT_FIVE:
            return 1.5
        if rule == WD_LINE_SPACING.DOUBLE:
            return 2.0
        # When rule is MULTIPLE, python-docx typically stores a float in line_spacing
        if rule == WD_LINE_SPACING.MULTIPLE and isinstance(ls, (int, float)):
            return float(ls)
        # If rule is None or AT_LEAST/EXACTLY we can't reliably map to a multiplier
        if isinstance(ls, (int, float)):
            return float(ls)
    except Exception:
        return None
    return None


class DocxSpacingPatternTool(Tool):
    name = "docx_spacing_pattern"
    description = (
        "Validate the line spacing pattern for the first three paragraphs: "
        "introduction=1.0, body=2.0, conclusion=1.5. "
        "IMPORTANT: file_path must be a Host path. If your file is inside the VM (e.g., /home/user/...), "
        "use get_vm_file(vm_path, dest_name) to download it to Host first and pass the returned Host path."
    )
    inputs = {
        "file_path": {
            "description": "Path to the .docx file on the Host (absolute or relative).",
            "type": "string",
        }
    }
    output_type = "string"

    def forward(self, file_path: str) -> str:
        return self.__call__(file_path)

    def __call__(self, file_path: str) -> str:
        try:
            if not file_path:
                return json.dumps({"passed": False, "details": "file_path is required"}, ensure_ascii=False)
            abs_path = os.path.abspath(file_path)
            if not os.path.exists(abs_path):
                return json.dumps({
                    "passed": False,
                    "details": (
                        f"File not found on host: {abs_path}. If your file is in the VM, use get_vm_file to download it first."
                    )
                }, ensure_ascii=False)

            doc = Document(abs_path)
            if len(doc.paragraphs) < 3:
                return json.dumps({
                    "passed": False,
                    "details": f"Document has only {len(doc.paragraphs)} paragraphs; need at least 3 to check spacing pattern"
                }, ensure_ascii=False)

            expected = [1.0, 2.0, 1.5]
            got = []
            undet_indexes = []
            for i in range(3):
                val = _get_spacing_value(doc.paragraphs[i])
                if val is None:
                    undet_indexes.append(i + 1)
                got.append(val)

            if undet_indexes:
                return json.dumps({
                    "passed": False,
                    "details": f"Could not determine line spacing for paragraphs: {undet_indexes}. Got: {got}"
                }, ensure_ascii=False)

            tol = 0.05
            comparisons = [abs(got[i] - expected[i]) <= tol for i in range(3)]
            passed = all(comparisons)
            if passed:
                return json.dumps({
                    "passed": True,
                    "details": f"Spacing pattern matches expected [1.0, 2.0, 1.5]; got {got}"
                }, ensure_ascii=False)
            else:
                return json.dumps({
                    "passed": False,
                    "details": f"Spacing pattern mismatch. Expected [1.0, 2.0, 1.5]; got {got}"
                }, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"passed": False, "details": f"Error: {e}"}, ensure_ascii=False)

    def to_code_prompt(self) -> str:
        return (
            "def docx_spacing_pattern(file_path: str) -> str:\n"
            "    '''Validate line spacing pattern: intro=1.0, body=2.0, conclusion=1.5 (first three paragraphs).\n"
            "    file_path must be a Host path; if the file is in the VM, first use get_vm_file.\n"
            "    Returns a JSON string with keys passed (bool) and details (str).'''\n"
        )
