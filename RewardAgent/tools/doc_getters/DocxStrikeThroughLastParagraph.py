from __future__ import annotations

import json
import os

from smolagents import Tool


class DocxStrikeThroughLastParagraphTool(Tool):
    name = "docx_strikethrough_last_paragraph"
    description = (
        "Verify that all runs in the last paragraph of the DOCX have strike-through formatting. "
        "This is a single-file adaptation of evaluate_strike_through_last_paragraph. "
        "IMPORTANT: file_path must be a Host path. If your file is inside the VM (e.g., /home/user/...), "
        "use get_vm_file(vm_path, dest_name) to download it to Host first and pass the returned Host path."
    )
    inputs = {
        "file_path": {
            "description": "Path to the .docx file on the Host (absolute or relative).",
            "type": "string",
        }
    }
    output_type = "string"

    def forward(self, file_path: str) -> str:
        return self.__call__(file_path)

    def __call__(self, file_path: str) -> str:
        try:
            if not file_path:
                return json.dumps({"passed": False, "details": "file_path is required"}, ensure_ascii=False)
            abs_path = os.path.abspath(file_path)
            if not os.path.exists(abs_path):
                return json.dumps({
                    "passed": False,
                    "details": (
                        f"File not found on host: {abs_path}. If your file is in the VM, use get_vm_file to download it first."
                    )
                }, ensure_ascii=False)

            try:
                from docx import Document
            except Exception as e:
                return json.dumps({"passed": False, "details": f"python-docx import error: {e}"}, ensure_ascii=False)

            try:
                document = Document(abs_path)
            except Exception as e:
                return json.dumps({"passed": False, "details": f"Failed to open DOCX: {e}"}, ensure_ascii=False)

            if not document.paragraphs:
                return json.dumps({"passed": False, "details": "Document has no paragraphs"}, ensure_ascii=False)

            last_paragraph = document.paragraphs[-1]
            runs = list(last_paragraph.runs)
            if len(runs) == 0:
                return json.dumps({"passed": False, "details": "Last paragraph has no runs"}, ensure_ascii=False)

            # Require every run with any non-empty text to be strike-through
            for run in runs:
                text = (run.text or "").strip()
                if text and not run.font.strike:
                    return json.dumps({
                        "passed": False,
                        "details": "Found non-strikethrough text in last paragraph"
                    }, ensure_ascii=False)

            return json.dumps({"passed": True, "details": "All text in last paragraph is strike-through"}, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"passed": False, "details": f"Error: {e}"}, ensure_ascii=False)

    def to_code_prompt(self) -> str:
        return (
            "def docx_strikethrough_last_paragraph(file_path: str) -> str:\n"
            "    '''Verify that all runs with text in the last paragraph have strike-through enabled.\n"
            "    file_path must be a Host path; if the file is in the VM, first use get_vm_file.\n"
            "    Returns a JSON string with keys passed (bool) and details (str).'''\n"
        )
