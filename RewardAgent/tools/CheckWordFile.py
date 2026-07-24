from __future__ import annotations

import json
import os
import re
import shlex
import subprocess
import sys
import tempfile
import time
from dataclasses import dataclass
from typing import Any, Dict, Optional, Tuple

from smolagents import Tool

API_URL = (os.getenv("IRA_BASE_URL") or os.getenv("OPENAI_BASE_URL") or "https://api.cometapi.com/v1/").rstrip("/") + "/"
MODEL_NAME = "o3"

GENERATED_CHECKER_PATH = os.path.join(tempfile.gettempdir(), f"ira_word_checker_{os.getpid()}.py")
HTTP_RETRY = 3
HTTP_TIMEOUT_SEC = 120

# ==== Auto-injected extraction helpers for DOCX (copied from word-test/exame.py, trimmed) ====
EXTRACT_PRELUDE = r"""
# ==== Auto-injected extraction helpers for DOCX (simplified from extractdocxinfo.py) ====
from typing import Any, Dict, List, Optional

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.enum.text import WD_ALIGN_PARAGRAPH

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"



def length_to_pt(value) -> Optional[float]:
    try:
        if value is None:
            return None
        return round(value.pt, 2)
    except Exception:
        return None



def rgb_color_to_hex(color) -> Optional[str]:
    try:
        if color is None:
            return None
        rgb = getattr(color, "rgb", None)
        if rgb is not None:
            return str(rgb)
        theme_color = getattr(color, "theme_color", None)
        if theme_color is not None:
            return f"THEME:{getattr(theme_color, 'name', str(theme_color))}"
    except Exception:
        pass
    return None



def extract_run_info(run) -> Dict[str, Any]:
    f = getattr(run, 'font', None)
    return {
        "text": getattr(run, 'text', ''),
        "font": {
            "name": getattr(f, 'name', None) if f is not None else None,
            "size_pt": length_to_pt(getattr(f, 'size', None) if f is not None else None),
            "bold": getattr(f, 'bold', None) if f is not None else None,
            "italic": getattr(f, 'italic', None) if f is not None else None,
            "underline": getattr(f, 'underline', None) if f is not None else None,
            "color_rgb": rgb_color_to_hex(getattr(f, 'color', None) if f is not None else None),
        },
    }



def _get_paragraph_alignment_name(p: Paragraph) -> Optional[str]:
    try:
        if p.alignment is None:
            return None
        return getattr(p.alignment, 'name', None) or str(p.alignment)
    except Exception:
        return None



def extract_paragraph_info(p: Paragraph, order_index: int) -> Dict[str, Any]:
    runs = []
    try:
        for r in getattr(p, 'runs', []):
            runs.append(extract_run_info(r))
    except Exception:
        pass
    return {
        "type": "paragraph",
        "order": order_index,
        "text": getattr(p, 'text', ''),
        "style_name": getattr(getattr(p, 'style', None), 'name', None),
        "alignment": _get_paragraph_alignment_name(p),
        "runs": runs,
    }



def extract_cell_paragraphs(cell: _Cell) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    try:
        for idx, para in enumerate(getattr(cell, 'paragraphs', [])):
            items.append(extract_paragraph_info(para, idx))
    except Exception:
        pass
    return items



def extract_table_info(tbl: Table, order_index: int) -> Dict[str, Any]:
    cells: List[Dict[str, Any]] = []
    rows = 0
    cols = 0
    try:
        rows = len(tbl.rows)
        cols = len(tbl.columns)
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                cells.append({
                    "row": r_idx,
                    "col": c_idx,
                    "text": getattr(cell, 'text', ''),
                    "paragraphs": extract_cell_paragraphs(cell),
                })
    except Exception:
        pass
    return {"type": "table", "order": order_index, "rows": rows, "cols": cols, "cells": cells}



def extract_body(doc: Document) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    try:
        for i, p in enumerate(getattr(doc, 'paragraphs', [])):
            items.append(extract_paragraph_info(p, i))
        for j, t in enumerate(getattr(doc, 'tables', [])):
            items.append(extract_table_info(t, j))
    except Exception:
        pass
    return items
# ==== End of helpers ====
"""


@dataclass
class VerifierResult:
    passed: bool
    reason: str
    llm_code_path: str
    stdout: str
    stderr: str

    def to_dict(self) -> Dict[str, Any]:
        return {
            "passed": self.passed,
            "reason": self.reason,
            "llm_code_path": self.llm_code_path,
            "stdout": self.stdout,
            "stderr": self.stderr,
        }


def _load_api_key() -> str:
    key = os.getenv("IRA_API_KEY") or os.getenv("OPENAI_API_KEY") or os.getenv("deerapi_key")
    if not key:
        try:
            from dotenv import load_dotenv
            load_dotenv()
            key = os.getenv("IRA_API_KEY") or os.getenv("OPENAI_API_KEY") or os.getenv("deerapi_key")
        except Exception:
            pass
    if not key:
        raise RuntimeError("Missing API key. Set IRA_API_KEY in the environment or a .env file.")
    return key


def _build_client() -> Any:
    api_key = _load_api_key()
    import openai  # type: ignore
    client = openai.OpenAI(api_key=api_key, base_url=API_URL)
    return client


def _call_llm_via_client(system: str, user: str) -> str:
    client = _build_client()
    resp = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        temperature=0,
    )
    content = getattr(resp.choices[0].message, "content", None)
    if isinstance(content, str):
        return content
    try:
        return resp.choices[0].message["content"]  # type: ignore[index]
    except Exception:
        return ""


def _call_llm_via_http(system: str, user: str) -> str:
    api_key = _load_api_key()
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    payload = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0,
    }
    last_err = None
    url = f"{API_URL}chat/completions"
    for i in range(1, HTTP_RETRY + 1):
        try:
            try:
                import requests  # type: ignore
                r = requests.post(url, headers=headers, json=payload, timeout=HTTP_TIMEOUT_SEC)
                r.raise_for_status()
                data = r.json()
            except Exception as e_req:
                from urllib.request import Request, urlopen
                import json as _json
                try:
                    req = Request(url, data=_json.dumps(payload).encode("utf-8"), headers=headers, method="POST")
                    with urlopen(req, timeout=HTTP_TIMEOUT_SEC) as resp:
                        txt = resp.read().decode("utf-8")
                    data = _json.loads(txt)
                except Exception as e_url:
                    last_err = e_url if e_req is None else e_req
                    raise last_err
            try:
                return data["choices"][0]["message"]["content"]
            except Exception as e:
                last_err = e
        except Exception as e:
            last_err = e
        if i < HTTP_RETRY:
            time.sleep(2 * i)
    if last_err:
        raise last_err
    return ""


def _llm_prompt(checks_text: str, error_context: Optional[str] = None) -> Tuple[str, str]:
    system = (
        "You are a careful Python auditor that writes a single, self-contained Python script "
        "to verify whether a set of simple checks have been satisfied on a DOCX file. "
        "Output only one Python code block fenced by triple backticks. Do not add any extra text. "
        "Your script must:")
    system += (
        "\n- Use python-docx (docx module) to read the DOCX file in read-only manner."
        "\n- Implement: def check(docx_path: str, checks: str) -> dict with keys: passed (bool), details (str)."
        "\n- Provide a __main__ that parses two required CLI args: --docx and --checks; then prints exactly one JSON line to stdout (no extra text)."
        "\n- Never modify files. No network. No os/subprocess/requests/shutil/socket/urllib usage. No eval/exec. Avoid sys.exit unless necessary."
        "\n- Treat the input 'checks' string as simple, atomic validation items (e.g., text contains X, paragraph style is Heading 1, table has N rows)."
        "\n- Choose which parts to inspect strictly from the provided checks text. If the checks do not specify a specific paragraph/table, you may check all and state your assumption in details."
        "\n- Be robust to text in paragraphs and text inside table cells; iterate runs and inspect formatting where relevant to the checks."
        "\n- If the required formatting properties are ambiguous, provide a conservative judgement and explain it in details."
        "\n- python-docx tips: for p in doc.paragraphs: for r in p.runs: ...; for tbl in doc.tables: for row in tbl.rows: for cell in row.cells: for p in cell.paragraphs: ..."
        "\n- Underline: Prefer boolean detection: treat underlined as (run.font.underline is True)."
        "\n- To read RGB from run.font.color.rgb, convert to string like str(rgb) -> 'RRGGBB' or compare tuple indices if available; rgb may be None for theme/auto."
        "\n- Your script must never crash; wrap main execution in try/except and always print one JSON line with keys 'passed' and 'details'."
    )

    system += (
        "\n\nNote: Predefined helper functions for robust DOCX extraction are injected at the top of the script and available to call: "
        "extract_body, extract_paragraph_info, extract_table_info, extract_run_info, rgb_color_to_hex. "
        "Use these helpers to focus on the check-specific verification logic."
    )

    system += (
        "\n\npython-docx Quick Reference:" \
        "\n- Imports:" \
        "\n  from docx import Document" \
        "\n- Paragraphs & runs:" \
        "\n  for p in doc.paragraphs: for r in p.runs: ...; p.alignment; p.style.name" \
        "\n- Tables:" \
        "\n  for tbl in doc.tables: for row in tbl.rows: for cell in row.cells: for p in cell.paragraphs: for r in p.runs: ..." \
        "\n- Fonts & colors:" \
        "\n  f = run.font; f.name; getattr(f.size, 'pt', None); f.bold; f.italic; f.underline; color = f.color; getattr(color, 'rgb', None)" \
        "\n- Output & safety:" \
        "\n  Print exactly one JSON line to stdout with keys 'passed' and 'details'. No file writes, no network, no os/subprocess/requests/shutil/socket/urllib, no eval/exec."
    )

    user = (
        "Write the script now. It must be a single code block:"\
        "\n```python\n# your script here\n```\n\n"\
        "The checks to perform are:\n" + checks_text + "\n"\
        "The script will be executed as: python word_checker.py --docx /abs/path.docx --checks '<CHECKS>'\n"\
        "Remember to print only a single JSON line on stdout with keys: passed, details."
    )
    if error_context:
        user += ("\n\nThe previous attempt failed with the following error/output. "
                 "Regenerate a corrected script that avoids these issues and adheres to the rules.\n" + error_context)
    return system, user


def _extract_code_block(text: str) -> Optional[str]:
    pattern = r"```(?:python)?\n([\s\S]*?)\n```"
    m = re.search(pattern, text, re.IGNORECASE)
    if not m:
        return None
    return m.group(1)


def _basic_static_safety_check(code: str) -> Optional[str]:
    forbidden_imports = ["os", "subprocess", "shutil", "socket", "requests", "urllib", "pptx"]
    for name in forbidden_imports:
        if re.search(rf"^\s*(import|from)\s+{re.escape(name)}\b", code, re.MULTILINE):
            return f"Forbidden import detected: {name}"
    if re.search(r"\beval\s*\(", code):
        return "Forbidden call: eval()"
    if re.search(r"\bexec\s*\(", code):
        return "Forbidden call: exec()"
    if re.search(r"open\s*\(.*['\"]\s*[wax]\s*['\"]", code):
        return "Forbidden file write mode detected"
    if re.search(r"http[s]?://", code):
        return "Forbidden network usage detected"
    return None


def _write_generated_checker(code: str, target_path: Optional[str] = None) -> str:
    path = target_path or GENERATED_CHECKER_PATH
    os.makedirs(os.path.dirname(path), exist_ok=True)
    prelude = EXTRACT_PRELUDE.strip() + "\n\n"
    with open(path, "w", encoding="utf-8") as f:
        f.write(prelude + code)
    return path


def _build_checker_cmd(checker_path: str, docx_path: str, checks_text: str) -> str:
    return f"{shlex.quote(sys.executable)} {shlex.quote(checker_path)} --docx {shlex.quote(docx_path)} --checks {shlex.quote(checks_text)}"


def _run_generated_checker(checker_path: str, docx_path: str, checks_text: str, timeout_sec: int = 20) -> Tuple[str, str, int, str]:
    cmd = _build_checker_cmd(checker_path, docx_path, checks_text)
    proc = subprocess.run(
        cmd,
        shell=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=timeout_sec,
        text=True,
    )
    return proc.stdout.strip(), proc.stderr.strip(), proc.returncode, cmd


def _parse_json_line(s: str) -> Optional[Dict[str, Any]]:
    line = s.strip().splitlines()[0] if s.strip() else ""
    candidate = line
    if not candidate or not candidate.strip().startswith("{"):
        m = re.search(r"\{[\s\S]*\}", s)
        candidate = m.group(0) if m else ""
    if not candidate:
        return None
    try:
        return json.loads(candidate)
    except Exception:
        return None


def _ensure_dir(path: str) -> None:
    try:
        os.makedirs(path, exist_ok=True)
    except Exception:
        pass


def _safe_write(path: str, content: str) -> None:
    try:
        _ensure_dir(os.path.dirname(path))
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
    except Exception:
        pass


def _log(trace: bool, msg: str) -> None:
    if trace:
        try:
            sys.stderr.write(msg.rstrip("\n") + "\n")
            sys.stderr.flush()
        except Exception:
            pass


def run_o3_verifier_word(checks_text: str, docx_path: str, attempts: int = 3, trace: bool = False, log_root_dir: Optional[str] = None, output_dir: Optional[str] = None) -> VerifierResult:
    if not checks_text:
        raise ValueError("checks_text is required")
    if not docx_path:
        raise ValueError("docx_path is required")
    if not os.path.isabs(docx_path):
        docx_path = os.path.abspath(docx_path)
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    ts = time.strftime("%Y%m%d-%H%M%S")
    if output_dir:
        base = os.path.join(output_dir, "doc_tools")
        log_root_dir = log_root_dir or os.path.join(base, "attempt_logs", f"word-{ts}")
        checker_path = os.path.join(base, "word_checker.py")
    else:
        if log_root_dir is None:
            log_root_dir = os.path.join(tempfile.gettempdir(), "ira_attempt_logs", f"word-{ts}")
        checker_path = GENERATED_CHECKER_PATH
    _ensure_dir(log_root_dir)

    summary_path = os.path.join(log_root_dir, "summary.txt")

    attempt = 1
    error_ctx: Optional[str] = None
    last_stdout = ""
    last_stderr = ""
    last_code_path = ""
    last_content = ""

    while attempt <= max(1, attempts):
        attempt_dir = os.path.join(log_root_dir, f"attempt-{attempt}")
        _ensure_dir(attempt_dir)
        _log(trace, f"[word attempt {attempt}] starting...")

        system, user = _llm_prompt(checks_text, error_context=error_ctx)
        _safe_write(os.path.join(attempt_dir, "00-prompt-system.txt"), system)
        _safe_write(os.path.join(attempt_dir, "01-prompt-user.txt"), user)

        content = ""
        err_sdk: Optional[Exception] = None
        used_path = "sdk"
        try:
            content = _call_llm_via_client(system, user)
        except Exception as e:
            err_sdk = e

        if not content or not _extract_code_block(content):
            used_path = "http"
            try:
                content = _call_llm_via_http(system, user)
            except Exception as e:
                meta = [
                    f"used_path={used_path}",
                    f"sdk_error={repr(err_sdk)}",
                    f"http_error={repr(e)}",
                ]
                _safe_write(os.path.join(attempt_dir, "02b-llm-meta.txt"), "\n".join(meta))
                reason = f"LLM call failed (sdk={err_sdk}) (http={e})"
                if attempt >= max(1, attempts):
                    _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                    _safe_write(summary_path, f"attempt {attempt}: LLM call failed. {reason}\n")
                    _log(trace, f"[word attempt {attempt}] LLM call failed. {reason}")
                    return VerifierResult(
                        passed=False,
                        reason=reason,
                        llm_code_path=last_code_path,
                        stdout=str(last_stdout or content),
                        stderr=str(last_stderr),
                    )
                else:
                    error_ctx = f"Previous LLM call failed. Reason: {reason}"
                    _safe_write(os.path.join(attempt_dir, "02-llm-raw.txt"), str(content))
                    _safe_write(os.path.join(attempt_dir, "02b-llm-meta.txt"), "\n".join(meta))
                    _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                    _safe_write(summary_path, f"attempt {attempt}: LLM call failed. Will retry.\n")
                    _log(trace, f"[word attempt {attempt}] LLM call failed, retrying...")
                    attempt += 1
                    continue

        last_content = content
        _safe_write(os.path.join(attempt_dir, "02-llm-raw.txt"), str(content))
        _safe_write(os.path.join(attempt_dir, "02b-llm-meta.txt"), f"used_path={used_path}\nsdk_error={repr(err_sdk)}")

        code = _extract_code_block(content)
        if not code:
            _safe_write(os.path.join(attempt_dir, "03-extracted_code.py"), code or "")
            if attempt >= max(1, attempts):
                _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                _safe_write(summary_path, f"attempt {attempt}: no valid code block.\n")
                _log(trace, f"[word attempt {attempt}] no valid code block; stop.")
                return VerifierResult(
                    passed=False,
                    reason="Failed to extract code block from LLM response",
                    llm_code_path=last_code_path,
                    stdout=content,
                    stderr="",
                )
            else:
                error_ctx = "The earlier response did not include a valid ```python code block. Ensure you wrap the entire script in triple backticks."
                _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                _safe_write(summary_path, f"attempt {attempt}: no code block; retry.\n")
                _log(trace, f"[word attempt {attempt}] no code block; retrying...")
                attempt += 1
                continue

        _safe_write(os.path.join(attempt_dir, "03-extracted_code.py"), code)

        unsafe_reason = _basic_static_safety_check(code)
        if unsafe_reason:
            _safe_write(os.path.join(attempt_dir, "04-safety_check.txt"), f"REJECTED: {unsafe_reason}")
            if attempt >= max(1, attempts):
                _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                _safe_write(summary_path, f"attempt {attempt}: safety check failed: {unsafe_reason}.\n")
                _log(trace, f"[word attempt {attempt}] safety check failed; stop.")
                return VerifierResult(
                    passed=False,
                    reason=f"Generated code rejected by safety check: {unsafe_reason}",
                    llm_code_path=last_code_path,
                    stdout=code,
                    stderr="",
                )
            else:
                error_ctx = f"Your previous code failed safety checks: {unsafe_reason}. Remove the offending imports/calls and regenerate."
                _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                _safe_write(summary_path, f"attempt {attempt}: safety check failed; retry.\n")
                _log(trace, f"[word attempt {attempt}] safety check failed; retrying...")
                attempt += 1
                continue
        else:
            _safe_write(os.path.join(attempt_dir, "04-safety_check.txt"), "OK")

        path = _write_generated_checker(code, target_path=checker_path)
        last_code_path = path
        per_attempt_copy = os.path.join(attempt_dir, f"generated_checker_attempt_{attempt}.py")
        _safe_write(per_attempt_copy, code)
        _safe_write(os.path.join(attempt_dir, "05-generated_checker_path.txt"), f"main={path}\ncopy={per_attempt_copy}")

        try:
            stdout, stderr, rc, cmd = _run_generated_checker(checker_path, docx_path, checks_text)
            _safe_write(os.path.join(attempt_dir, "06-run_command.txt"), cmd)
            _safe_write(os.path.join(attempt_dir, "07-stdout.txt"), stdout)
            _safe_write(os.path.join(attempt_dir, "08-stderr.txt"), stderr)
        except subprocess.TimeoutExpired:
            _safe_write(os.path.join(attempt_dir, "06-run_command.txt"), _build_checker_cmd(checker_path, docx_path, checks_text))
            _safe_write(os.path.join(attempt_dir, "08-stderr.txt"), "TimeoutExpired")
            if attempt >= max(1, attempts):
                _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                _safe_write(summary_path, f"attempt {attempt}: checker timed out.\n")
                _log(trace, f"[word attempt {attempt}] checker timed out; stop.")
                return VerifierResult(
                    passed=False,
                    reason="Generated checker timed out",
                    llm_code_path=path,
                    stdout=last_stdout,
                    stderr=last_stderr,
                )
            else:
                error_ctx = "Your previous script timed out. Make the logic efficient and avoid long loops."
                _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                _safe_write(summary_path, f"attempt {attempt}: checker timed out; retry.\n")
                _log(trace, f"[word attempt {attempt}] checker timed out; retrying...")
                attempt += 1
                continue
        except Exception as e:
            _safe_write(os.path.join(attempt_dir, "06-run_command.txt"), _build_checker_cmd(checker_path, docx_path, checks_text))
            _safe_write(os.path.join(attempt_dir, "08-stderr.txt"), f"Exception: {repr(e)}")
            if attempt >= max(1, attempts):
                _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                _safe_write(summary_path, f"attempt {attempt}: failed to run checker: {repr(e)}.\n")
                _log(trace, f"[word attempt {attempt}] failed to run checker; stop.")
                return VerifierResult(
                    passed=False,
                    reason=f"Failed to run generated checker: {e}",
                    llm_code_path=path,
                    stdout=last_stdout,
                    stderr=last_stderr,
                )
            else:
                error_ctx = f"Your previous script raised an exception on execution: {e}. Fix and regenerate."
                _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                _safe_write(summary_path, f"attempt {attempt}: checker exception; retry.\n")
                _log(trace, f"[word attempt {attempt}] checker exception; retrying...")
                attempt += 1
                continue

        last_stdout, last_stderr = stdout, stderr
        parsed = _parse_json_line(stdout)
        if parsed:
            _safe_write(os.path.join(attempt_dir, "09-parsed_json.json"), json.dumps(parsed, ensure_ascii=False, indent=2))
        else:
            _safe_write(os.path.join(attempt_dir, "09-parsed_json.json"), "<invalid or missing JSON line>")

        if not parsed or not isinstance(parsed, dict) or "passed" not in parsed:
            short_err = (stderr or "").strip()
            if len(short_err) > 800:
                short_err = short_err[-800:]
            if attempt >= max(1, attempts):
                _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                _safe_write(summary_path, f"attempt {attempt}: invalid JSON output.\n")
                _log(trace, f"[word attempt {attempt}] invalid JSON output; stop.")
                return VerifierResult(
                    passed=False,
                    reason="Generated checker did not produce a valid single-line JSON with 'passed'",
                    llm_code_path=path,
                    stdout=stdout,
                    stderr=stderr,
                )
            else:
                error_ctx = (
                    "Your script must print exactly one JSON line on stdout with keys 'passed' and 'details'. "
                    "The previous run failed. Here is stderr/issue summary:\n" + short_err + "\n\n"
                    "Also wrap your main in try/except so exceptions still result in a single JSON line."
                )
                _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                _safe_write(summary_path, f"attempt {attempt}: invalid JSON output; retry.\n")
                _log(trace, f"[word attempt {attempt}] invalid JSON; retrying...")
                attempt += 1
                continue

        passed = bool(parsed.get("passed"))
        details = str(parsed.get("details", ""))

        if not passed:
            short_details = details.strip()
            if len(short_details) > 800:
                short_details = short_details[:800]
            if attempt >= max(1, attempts):
                _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                _safe_write(summary_path, f"attempt {attempt}: checker returned failed.\n")
                _log(trace, f"[word attempt {attempt}] checker returned failed; stop.")
                return VerifierResult(
                    passed=False,
                    reason=details,
                    llm_code_path=path,
                    stdout=stdout,
                    stderr=stderr,
                )
            else:
                error_ctx = (
                    "Your previous checker returned passed=false with details: " + short_details + "\n"
                    "Regenerate a corrected checker that fixes this issue while strictly following all rules."
                )
                _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
                _safe_write(summary_path, f"attempt {attempt}: checker returned failed; retry.\n")
                _log(trace, f"[word attempt {attempt}] checker returned failed; retrying...")
                last_stdout, last_stderr = stdout, stderr
                attempt += 1
                continue

        _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
        _safe_write(summary_path, f"attempt {attempt}: SUCCESS.\n")
        _log(trace, f"[word attempt {attempt}] SUCCESS. Artifacts at: {attempt_dir}")
        return VerifierResult(
            passed=True,
            reason=details,
            llm_code_path=path,
            stdout=stdout,
            stderr=stderr,
        )

    _safe_write(summary_path, (open(summary_path, "r", encoding="utf-8").read() + "\n") if os.path.exists(summary_path) else "")
    _safe_write(summary_path, "No attempt yielded a valid result.\n")
    _log(trace, f"All attempts exhausted. Logs at: {log_root_dir}")
    return VerifierResult(
        passed=False,
        reason="Exhausted attempts without a valid result",
        llm_code_path=last_code_path,
        stdout=last_stdout or last_content,
        stderr=last_stderr,
    )


class CheckWordFileTool(Tool):
    name = "checkwordfile"
    description = (
        "Verify a Word (.docx) file content/format according to simple check items (things_to_check). "
        "This tool requires a Host file path. If your file is inside the VM (e.g., /home/user/...), "
        "please first use get_vm_file(vm_path, dest_name) to download it to the Host and then pass the returned Host path here. "
        "Returns a JSON string with keys passed, reason, llm_code_path, stdout, stderr."
    )
    inputs = {
        "file_path": {
            "description": (
                "Path to the .docx file on the Host (absolute or relative). "
                "Do not pass VM paths. If the file is in the VM, call get_vm_file first and then pass the returned Host path."
            ),
            "type": "string",
        },
        "things_to_check": {
            "description": "Simple text describing the items to verify in the Word file (e.g., 'Title is bold; body contains \"Company\"').",
            "type": "string",
        },
    }
    output_type = "string"

    def __init__(self, output_dir: Optional[str] = None):
        super().__init__()
        self.output_dir = output_dir

    def set_output_dir(self, output_dir: str) -> None:
        self.output_dir = output_dir

    def forward(self, file_path: str, things_to_check: str) -> str:
        return self.__call__(file_path, things_to_check)

    def __call__(self, file_path: str, things_to_check: str) -> str:
        try:
            if not file_path:
                return json.dumps({"passed": False, "reason": "file_path is required"}, ensure_ascii=False)
            abs_path = os.path.abspath(file_path)
            if not os.path.exists(abs_path):
                return json.dumps({
                    "passed": False,
                    "reason": (
                        f"File not found on host: {abs_path}. "
                        "This tool requires a Host file path. If your file is inside the VM (e.g., /home/user/...), "
                        "please use get_vm_file to download it to the Host and pass the returned Host path."
                    )
                }, ensure_ascii=False)
            if not things_to_check:
                return json.dumps({"passed": False, "reason": "things_to_check is required"}, ensure_ascii=False)

            res = run_o3_verifier_word(things_to_check, abs_path, attempts=3, trace=False, log_root_dir=None, output_dir=self.output_dir)
            return json.dumps(res.to_dict(), ensure_ascii=False)
        except Exception as e:
            return json.dumps({"passed": False, "reason": f"Error: {e}"}, ensure_ascii=False)

    def to_code_prompt(self) -> str:
        return (
            "def checkwordfile(file_path: str, things_to_check: str) -> str:\n"
            "    '''Verify a Word (.docx) file content/format according to simple check items (things_to_check).\n"
            "    Note: file_path must be a Host path; if the file is in the VM, first use get_vm_file to download it.\n"
            "    Returns a JSON string with keys passed, reason, llm_code_path, stdout, stderr.'''\n"
        )
